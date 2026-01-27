"""
PDF 工具模块业务逻辑
实现 PDF 阅读、编辑、转换等操作
"""

import os
import logging
import fitz  # PyMuPDF
import zipfile
import io
from typing import Optional, List, Tuple, Dict, Any
from sqlalchemy import select, func, and_
from sqlalchemy.ext.asyncio import AsyncSession
from pathlib import Path

from .pdf_models import Pdf
from .pdf_schemas import (
    PdfMetadata, 
    PdfMergeRequest, 
    PdfSplitRequest,
    PdfCompressRequest,
    PdfWatermarkRequest,
    PdfImagesToPdfRequest,
    PdfToImagesRequest,
    PdfToWordRequest,
    PdfToExcelRequest,
    PdfRemoveWatermarkRequest,
    PdfEncryptRequest,
    PdfDecryptRequest,
    PdfRotateRequest,
    PdfReorderRequest,
    PdfExtractPagesRequest,
    PdfAddPageNumbersRequest,
    PdfSignRequest,
    PdfWordToPdfRequest,
    PdfExcelToPdfRequest,
    PdfSaveTextRequest
)
from models.storage import FileRecord
from utils.storage import get_storage_manager
from utils.timezone import get_beijing_time

logger = logging.getLogger(__name__)


class PdfService:
    """
    PDF 工具服务类
    
    providing PDF reading, splitting, merging, conversion, and syncing with FileManager
    """

    @staticmethod
    async def get_file_path_by_storage(user_id: int, storage_path: str) -> Tuple[str, str]:
        """
        通过存储路径获取文件物理路径 (用于虚拟映射文件)
        """
        storage = get_storage_manager()
        file_path = storage.get_file_path(storage_path)
        if not file_path or not os.path.exists(file_path):
            raise ValueError(f"文件不存在: {storage_path}")
        return str(file_path), file_path.name

    @staticmethod
    async def get_file_path(db: AsyncSession, file_id: Optional[int], user_id: int, source: str = "filemanager") -> Tuple[str, str]:
        """
        统一获取文件物理路径和文件名
        source 支持: filemanager, upload
        """
        if file_id is None:
             raise ValueError("无效的文件ID")

        storage_path = None
        filename = None
        
        if source == "filemanager":
            # 动态导入，避免循环引用
            from modules.filemanager.filemanager_models import VirtualFile
            result = await db.execute(select(VirtualFile).where(and_(VirtualFile.id == file_id, VirtualFile.user_id == user_id)))
            virtual_file = result.scalar_one_or_none()
            if not virtual_file:
                raise ValueError(f"文件不存在或无权访问 (ID: {file_id})")
            storage_path = virtual_file.storage_path
            filename = virtual_file.name
        else:
            result = await db.execute(select(FileRecord).where(and_(FileRecord.id == file_id, FileRecord.uploader_id == user_id)))
            file_record = result.scalar_one_or_none()
            if not file_record:
                raise ValueError(f"文件不存在或无权访问 (ID: {file_id})")
            storage_path = file_record.storage_path
            filename = file_record.filename
            
        storage = get_storage_manager()
        file_path = storage.get_file_path(storage_path)
        if not file_path or not os.path.exists(file_path):
            raise ValueError(f"物理文件丢失: {filename}")
            
        return str(file_path), filename

    @staticmethod
    async def _resolve_file(db: AsyncSession, user_id: int, file_id: Optional[int] = None, path: Optional[str] = None) -> Tuple[str, str]:
        """统一解析请求中的 file_id 或 path，返回 (物理路径, 文件名)"""
        if file_id:
            # 默认先尝试从文件管理器找，找不到再尝试从 upload 记录找
            try:
                return await PdfService.get_file_path(db, file_id, user_id, "filemanager")
            except:
                return await PdfService.get_file_path(db, file_id, user_id, "upload")
        elif path:
            return await PdfService.get_file_path_by_storage(user_id, path)
        else:
            raise ValueError("未提供有效的文件 ID 或路径")

    @staticmethod
    def get_file_path_by_name(user_id: int, filename: str, category: str = "uploads") -> Tuple[str, str]:
        """
        通过文件名获取 PDF 模块存储的文件路径
        用于从 PDF 模块自己的存储目录获取文件
        """
        storage = get_storage_manager()
        # 使用更规范的方法获取目录
        pdf_dir = storage.get_module_dir("pdf", category, user_id)
        file_path = pdf_dir / filename
        
        # 安全检查：确保文件在正确目录内
        try:
            file_path.resolve().relative_to(pdf_dir.resolve())
        except ValueError:
            raise ValueError("非法的文件路径")
        
        if not file_path.exists():
            raise ValueError(f"文件不存在: {filename}")
            
        return str(file_path), filename

    @staticmethod
    def _get_output_path(user_id: int, output_name: str) -> Path:
        """获取输出路径（使用 outputs 子目录）"""
        storage = get_storage_manager()
        # 使用规范方法获取目录
        outputs_dir = storage.get_module_dir("pdf", "outputs", user_id)
        return outputs_dir / output_name

    @staticmethod
    async def _save_history(db: AsyncSession, user_id: int, title: str, filename: str, operation: str):
        """记录操作历史"""
        history = Pdf(
            user_id=user_id,
            title=title,
            filename=filename,
            operation=operation
        )
        db.add(history)

    @staticmethod
    async def get_metadata(file_path: str) -> PdfMetadata:
        """获取 PDF 元数据"""
        try:
            doc = fitz.open(file_path)
            # PyMuPDF metadata 可能为 None（例如加密文档未解锁时）
            meta = doc.metadata or {}
            
            # 即使加密，doc.page_count 通常也是可用的
            page_count = doc.page_count
            size_bytes = os.path.getsize(file_path)
            
            result = PdfMetadata(
                title=meta.get("title"),
                author=meta.get("author"),
                subject=meta.get("subject"),
                keywords=meta.get("keywords"),
                creator=meta.get("creator"),
                producer=meta.get("producer"),
                creationDate=meta.get("creationDate"),
                modDate=meta.get("modDate"),
                page_count=page_count,
                size_bytes=size_bytes,
                is_encrypted=doc.is_encrypted
            )
            doc.close()
            return result
        except Exception as e:
            logger.error(f"解析 PDF 元数据失败: {e}")
            # 如果彻底失败，返回一个提示加密或错误的对象
            return PdfMetadata(
                title="无法读取",
                page_count=0,
                is_encrypted=True # 假设它可能是加密导致的
            )


    @staticmethod
    async def render_page(file_path: str, page_num: int, zoom: float = 2.0) -> bytes:
        """渲染 PDF 指定页面为图片 (PNG)"""
        try:
            doc = fitz.open(file_path)
            if page_num < 0 or page_num >= doc.page_count:
                raise ValueError(f"页码超出范围 (0-{doc.page_count-1})")
            
            page = doc.load_page(page_num)
            mat = fitz.Matrix(zoom, zoom)
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img_data = pix.tobytes("png")
            doc.close()
            return img_data
        except Exception as e:
            logger.error(f"渲染 PDF 页面失败: {e}")
            raise ValueError(f"渲染失败: {str(e)}")

    @staticmethod
    async def merge_pdfs(db: AsyncSession, user_id: int, request: PdfMergeRequest) -> str:
        """合并多个 PDF"""
        try:
            result_doc = fitz.open()
            # 收集文件路径
            file_paths = []
            if request.file_ids:
                for fid in request.file_ids:
                    fpath, _ = await PdfService.get_file_path(db, fid, user_id)
                    file_paths.append(fpath)
            if request.paths:
                for p in request.paths:
                    fpath, _ = await PdfService.get_file_path_by_storage(user_id, p)
                    file_paths.append(fpath)

            for fpath in file_paths:
                with fitz.open(fpath) as doc:
                    result_doc.insert_pdf(doc)
            
            output_name = request.output_name
            if not output_name.lower().endswith(".pdf"):
                output_name += ".pdf"
            
            save_path = PdfService._get_output_path(user_id, output_name)
            result_doc.save(str(save_path))
            result_doc.close()
            
            await PdfService._save_history(db, user_id, f"合并 PDF: {output_name}", output_name, "merge")

            
            return str(save_path)
        except Exception as e:
            logger.error(f"合并 PDF 失败: {e}")
            raise ValueError(f"合并失败: {str(e)}")

    @staticmethod
    async def split_pdf(db: AsyncSession, user_id: int, request: PdfSplitRequest) -> str:
        """拆分 PDF"""
        try:
            fpath, old_name = await PdfService._resolve_file(db, user_id, request.file_id, request.path)
            doc = fitz.open(fpath)
            
            pages_to_keep = []
            for part in request.page_ranges.split(','):
                part = part.strip()
                if not part: continue
                if '-' in part:
                    try:
                        start, end = map(int, part.split('-'))
                        pages_to_keep.extend(range(start - 1, end))
                    except ValueError:
                        continue
                else:
                    try:
                        pages_to_keep.append(int(part) - 1)
                    except ValueError:
                        continue
            
            pages_to_keep = sorted(list(set([p for p in pages_to_keep if 0 <= p < doc.page_count])))
            
            if not pages_to_keep:
                raise ValueError("未选择有效的页码范围")
            
            # 使用 select 方法保留/重排选中的页面
            # 注意：select 是就地修改，但只要我们不覆盖源文件保存，就是安全的
            doc.select(pages_to_keep)
            
            output_name = request.output_name or f"split_{old_name}"
            if not output_name.lower().endswith(".pdf"):
                output_name += ".pdf"
                
            save_path = PdfService._get_output_path(user_id, output_name)
            doc.save(str(save_path))
            doc.close()
            
            await PdfService._save_history(db, user_id, f"拆分 PDF: {output_name}", output_name, "split")

            
            return str(save_path)
        except Exception as e:
            logger.error(f"拆分 PDF 失败: {e}")
            raise ValueError(f"拆分失败: {str(e)}")

    @staticmethod
    async def extract_text(file_path: str) -> str:
        """提取 PDF 文本"""
        try:
            doc = fitz.open(file_path)
            text = ""
            for page in doc:
                text += page.get_text() + "\n"
            doc.close()
            return text
        except Exception as e:
            logger.error(f"提取 PDF 文本失败: {e}")
            raise ValueError(f"文本提取失败: {str(e)}")

    @staticmethod
    async def compress_pdf(db: AsyncSession, user_id: int, request: PdfCompressRequest) -> str:
        """压缩 PDF"""
        try:
            # 延迟导入以避免循环依赖或不必要的加载
            from PIL import Image
            import io
            
            fpath, old_name = await PdfService._resolve_file(db, user_id, request.file_id, request.path)
            doc = fitz.open(fpath)
            
            # 只有当 level >= 2 时才启用激进的图片压缩
            # level 1: 仅清理 (Quick)
            # level 2: 图片压缩 (Standard)
            # level 3: 图片压缩+缩放 (Max)
            if request.level >= 2:
                processed_xrefs = set()
                
                for page in doc:
                    images = page.get_images()
                    for img in images:
                        xref = img[0]
                        if xref in processed_xrefs:
                            continue
                        processed_xrefs.add(xref)
                        
                        try:
                            # 提取图片
                            pix = fitz.Pixmap(doc, xref)
                            
                            # 如果图片很小 (<50KB) 或尺寸很小，跳过
                            if pix.size < 50 * 1024 or pix.width < 100 or pix.height < 100:
                                pix = None
                                continue
                                
                            # 处理颜色空间
                            if pix.n - pix.alpha > 3: # CMYK 等
                                pix0 = fitz.Pixmap(fitz.csRGB, pix)
                                pix = pix0
                            
                            # 转 PIL 处理
                            img_data = pix.tobytes()
                            image = Image.open(io.BytesIO(img_data))
                            
                            # 确定压缩参数
                            quality = 75
                            new_size = None
                            
                            if request.level >= 3:
                                quality = 50
                                # 如果图片过大，进行缩放
                                if image.width > 1500 or image.height > 1500:
                                    factor = 1500 / max(image.width, image.height)
                                    new_size = (int(image.width * factor), int(image.height * factor))
                            
                            if new_size:
                                image = image.resize(new_size, Image.Resampling.LANCZOS)
                                
                            # 压缩
                            buffer = io.BytesIO()
                            # 统一转 RGB 并存为 JPEG
                            image.convert("RGB").save(buffer, format="JPEG", quality=quality, optimize=True)
                            new_bytes = buffer.getvalue()
                            
                            # 只有当压缩后体积确实减小时才替换
                            if len(new_bytes) < pix.size:
                                doc.update_stream(xref, new_bytes)
                                
                            pix = None
                        except Exception as e:
                            # 个别图片处理失败不应中断整个流程
                            logger.warning(f"压缩图片失败 xref={xref}: {e}")
                            continue

            output_name = request.output_name or f"compressed_{old_name}"
            if not output_name.lower().endswith(".pdf"):
                output_name += ".pdf"
            
            save_path = PdfService._get_output_path(user_id, output_name)
            
            # 使用最大程度的垃圾清理
            # garbage=4: dedup contents, streams, unused objects
            doc.save(str(save_path), garbage=4, deflate=True)
            doc.close()
            
            await PdfService._save_history(db, user_id, f"压缩 PDF: {output_name}", output_name, "compress")

            return str(save_path)
        except Exception as e:
            logger.error(f"压缩 PDF 失败: {e}")
            raise ValueError(f"压缩失败: {str(e)}")

    @staticmethod
    async def add_watermark(db: AsyncSession, user_id: int, request: PdfWatermarkRequest) -> str:
        """添加水印"""
        try:
            fpath, old_name = await PdfService._resolve_file(db, user_id, request.file_id, request.path)
            doc = fitz.open(fpath)
            
            output_name = request.output_name or f"watermarked_{old_name}"
            if not output_name.lower().endswith(".pdf"):
                output_name += ".pdf"
            
            for page in doc:
                rect = page.rect
                
                # 使用 insert_text 配合 morph 矩阵实现旋转
                # 计算页面中心
                center_point = fitz.Point(rect.width / 2, rect.height / 2)
                
                # 创建旋转矩阵 (顺时针 45 度)
                mat = fitz.Matrix(45)
                
                page.insert_text(
                    center_point,
                    request.text,
                    fontsize=request.fontsize,
                    fontname="china-s",  # 支持简体中文的内置字体
                    color=request.color,
                    fill_opacity=request.opacity,
                    morph=(center_point, mat)  # (固定点, 矩阵) 实现围绕该点旋转
                )
                
            save_path = PdfService._get_output_path(user_id, output_name)
            doc.save(str(save_path))
            doc.close()
            
            await PdfService._save_history(db, user_id, f"添加水印: {output_name}", output_name, "watermark")

            return str(save_path)
        except Exception as e:
            logger.error(f"添加水印失败: {e}")
            raise ValueError(f"添加水印失败: {str(e)}")

    @staticmethod
    async def images_to_pdf(db: AsyncSession, user_id: int, request: PdfImagesToPdfRequest) -> str:
        """图片转 PDF"""
        try:
            result_doc = fitz.open()
            
            # 收集文件路径
            file_paths = []
            if request.file_ids:
                for fid in request.file_ids:
                    fpath, _ = await PdfService.get_file_path(db, fid, user_id)
                    file_paths.append(fpath)
            if request.paths:
                for p in request.paths:
                    fpath, _ = await PdfService.get_file_path_by_storage(user_id, p)
                    file_paths.append(fpath)

            for fpath in file_paths:
                # 打开图片并转换为 PDF 单页
                img = fitz.open(fpath)
                pdfbytes = img.convert_to_pdf()
                img.close()
                img_pdf = fitz.open("pdf", pdfbytes)
                result_doc.insert_pdf(img_pdf)
            
            output_name = request.output_name
            if not output_name.lower().endswith(".pdf"):
                output_name += ".pdf"
                
            save_path = PdfService._get_output_path(user_id, output_name)
            result_doc.save(str(save_path))
            result_doc.close()
            
            await PdfService._save_history(db, user_id, f"图片转 PDF: {output_name}", output_name, "img2pdf")

            return str(save_path)
        except Exception as e:
            logger.error(f"图片转PDF失败: {e}")
            raise ValueError(f"转换失败: {str(e)}")

    @staticmethod
    async def pdf_to_images(db: AsyncSession, user_id: int, request: PdfToImagesRequest) -> str:
        """PDF 转图片 (Zip)"""
        try:
            fpath, old_name = await PdfService._resolve_file(db, user_id, request.file_id, request.path)
            doc = fitz.open(fpath)
            
            base_name = os.path.splitext(old_name)[0]
            zip_name = f"{base_name}_images.zip"
            save_path = PdfService._get_output_path(user_id, zip_name)
            
            with zipfile.ZipFile(save_path, 'w', zipfile.ZIP_DEFLATED) as zf:
                pages_to_process = range(doc.page_count)
                # 如果有范围过滤逻辑可以加在这里
                
                for i in pages_to_process:
                    page = doc.load_page(i)
                    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2)) # 2x zoom for quality
                    img_data = pix.tobytes("png")
                    zf.writestr(f"page_{i+1:03d}.png", img_data)
            
            doc.close()
            
            await PdfService._save_history(db, user_id, f"PDF 转图片: {zip_name}", zip_name, "pdf2img")
            await PdfService._register_to_filemanager(db, user_id, zip_name, str(save_path), os.path.getsize(save_path))
            return str(save_path)
        except Exception as e:
            logger.error(f"PDF转图片失败: {e}")
            raise ValueError(f"转换失败: {str(e)}")

    @staticmethod
    async def list_history(db: AsyncSession, user_id: int, page: int = 1, page_size: int = 20) -> Tuple[List[Pdf], int]:
        """获取 PDF 操作历史"""
        stmt = select(Pdf).where(Pdf.user_id == user_id).order_by(Pdf.created_at.desc())
        
        # 统计总数
        count_stmt = select(func.count()).select_from(stmt.subquery())
        total = (await db.execute(count_stmt)).scalar() or 0
        
        # 分页
        stmt = stmt.offset((page - 1) * page_size).limit(page_size)
        items = (await db.execute(stmt)).scalars().all()
        
        return list(items), total

    @staticmethod
    def _parse_page_ranges(page_ranges: str, total_pages: int) -> List[int]:
        """解析页码范围字符串，返回页码列表（0-indexed）"""
        pages = []
        for part in page_ranges.split(','):
            part = part.strip()
            if not part:
                continue
            if '-' in part:
                try:
                    start, end = map(int, part.split('-'))
                    pages.extend(range(start - 1, end))
                except ValueError:
                    continue
            else:
                try:
                    pages.append(int(part) - 1)
                except ValueError:
                    continue
        # 过滤有效页码并去重
        return sorted(list(set([p for p in pages if 0 <= p < total_pages])))

    @staticmethod
    async def pdf_to_word(db: AsyncSession, user_id: int, request: PdfToWordRequest) -> str:
        """PDF 转 Word 文档"""
        try:
            fpath, old_name = await PdfService._resolve_file(db, user_id, request.file_id, request.path)
            
            output_name = request.output_name or os.path.splitext(old_name)[0] + ".docx"
            if not output_name.lower().endswith(".docx"):
                output_name += ".docx"
            
            save_path = PdfService._get_output_path(user_id, output_name)
            
            # 尝试使用 pdf2docx 库（如果可用）
            try:
                from pdf2docx import Converter
                cv = Converter(fpath)
                cv.convert(str(save_path))
                cv.close()
            except ImportError:
                # 如果没有 pdf2docx，则提取文本并创建简单的 docx
                from docx import Document
                doc = fitz.open(fpath)
                word_doc = Document()
                
                for page_num in range(doc.page_count):
                    page = doc.load_page(page_num)
                    text = page.get_text()
                    if page_num > 0:
                        word_doc.add_page_break()
                    word_doc.add_paragraph(text)
                
                doc.close()
                word_doc.save(str(save_path))
            
            await PdfService._save_history(db, user_id, f"PDF 转 Word: {output_name}", output_name, "pdf2word")

            return str(save_path)
        except Exception as e:
            logger.error(f"PDF转Word失败: {e}")
            raise ValueError(f"转换失败: {str(e)}")

    @staticmethod
    async def pdf_to_excel(db: AsyncSession, user_id: int, request: PdfToExcelRequest) -> str:
        """PDF 转 Excel（提取表格）"""
        try:
            fpath, old_name = await PdfService._resolve_file(db, user_id, request.file_id, request.path)
            
            output_name = request.output_name or os.path.splitext(old_name)[0] + ".xlsx"
            if not output_name.lower().endswith(".xlsx"):
                output_name += ".xlsx"
            
            save_path = PdfService._get_output_path(user_id, output_name)
            
            doc = fitz.open(fpath)
            
            # 确定要处理的页面
            if request.page_ranges:
                pages_to_process = PdfService._parse_page_ranges(request.page_ranges, doc.page_count)
            else:
                pages_to_process = list(range(doc.page_count))
            
            # 使用 openpyxl 创建 Excel
            from openpyxl import Workbook
            wb = Workbook()
            ws = wb.active
            ws.title = "表格数据"
            
            row_idx = 1
            for page_num in pages_to_process:
                page = doc.load_page(page_num)
                # 提取表格
                # PyMuPDF 的 find_tables 返回 TableFinder 对象
                finder = page.find_tables()
                
                if finder.tables:
                    for table in finder:
                        # 写入表头分隔
                        ws.cell(row=row_idx, column=1, value=f"--- 第 {page_num + 1} 页表格 ---")
                        row_idx += 1
                        
                        # 提取表格数据
                        for row_data in table.extract():
                            for col_idx, cell_value in enumerate(row_data, 1):
                                ws.cell(row=row_idx, column=col_idx, value=cell_value or "")
                            row_idx += 1
                        row_idx += 1  # 表格之间空一行
                else:
                    # Fallback: 如果未检测到表格，提取纯文本
                    text = page.get_text()
                    if text.strip():
                        ws.cell(row=row_idx, column=1, value=f"--- 第 {page_num + 1} 页文本（未检测到表格）---")
                        row_idx += 1
                        for line in text.split('\n'):
                            if line.strip():
                                ws.cell(row=row_idx, column=1, value=line.strip())
                                row_idx += 1
                        row_idx += 1
            
            doc.close()
            wb.save(str(save_path))
            
            await PdfService._save_history(db, user_id, f"PDF 转 Excel: {output_name}", output_name, "pdf2excel")

            return str(save_path)
        except Exception as e:
            logger.error(f"PDF转Excel失败: {e}")
            raise ValueError(f"转换失败: {str(e)}")

    @staticmethod
    async def remove_watermark(db: AsyncSession, user_id: int, request: PdfRemoveWatermarkRequest) -> str:
        """去除 PDF 水印"""
        try:
            fpath, old_name = await PdfService._resolve_file(db, user_id, request.file_id, request.path)
            
            output_name = request.output_name or f"no_watermark_{old_name}"
            if not output_name.lower().endswith(".pdf"):
                output_name += ".pdf"
            
            doc = fitz.open(fpath)
            
            for page in doc:
                # 策略 1: 移除倾斜的文本 (常见水印特征)
                # 获取页面详细文本结构
                blocks = page.get_text("dict")["blocks"]
                for block in blocks:
                    if block["type"] == 0:  # 0 = 文本块
                        for line in block["lines"]:
                            # line["dir"] 是方向向量 (cos, sin)
                            # 水平文本: dir=(1, 0), sin~0
                            # 垂直文本: dir=(0, -1) 或 (0, 1), cos~0
                            # 如果既不水平也不垂直，认为是倾斜水印
                            dir_vec = line["dir"]
                            is_horizontal = abs(dir_vec[1]) < 0.05
                            is_vertical = abs(dir_vec[0]) < 0.05
                            
                            if not is_horizontal and not is_vertical:
                                # 添加遮盖注释（后续统一应用）
                                page.add_redact_annot(line["bbox"])

                # 策略 2: 移除可能是水印的大型半透明图形/图像
                # ... (保留原有或者简化原有逻辑)
                # 原有的基于图画的检测逻辑比较脆弱，我们主要依赖上面的文本检测
                
                # 应用所有遮盖
                page.apply_redactions()
            
            save_path = PdfService._get_output_path(user_id, output_name)
            doc.save(str(save_path), garbage=4, deflate=True)
            doc.close()
            
            await PdfService._save_history(db, user_id, f"去除水印: {output_name}", output_name, "remove_watermark")

            return str(save_path)
        except Exception as e:
            logger.error(f"去除水印失败: {e}")
            raise ValueError(f"去除水印失败: {str(e)}")

    @staticmethod
    async def encrypt_pdf(db: AsyncSession, user_id: int, request: PdfEncryptRequest) -> str:
        """加密 PDF"""
        try:
            fpath, old_name = await PdfService._resolve_file(db, user_id, request.file_id, request.path)
            
            output_name = request.output_name or f"encrypted_{old_name}"
            if not output_name.lower().endswith(".pdf"):
                output_name += ".pdf"
            
            doc = fitz.open(fpath)
            
            # 设置加密参数
            # 用户密码用于打开文档，所有者密码用于修改权限
            owner_pwd = request.owner_password or request.password
            
            perm = (
                fitz.PDF_PERM_ACCESSIBILITY |  # 允许辅助功能
                fitz.PDF_PERM_PRINT |          # 允许打印
                fitz.PDF_PERM_COPY             # 允许复制
            )
            
            save_path = PdfService._get_output_path(user_id, output_name)
            
            # 使用 AES-256 加密
            doc.save(
                str(save_path),
                encryption=fitz.PDF_ENCRYPT_AES_256,
                user_pw=request.password,
                owner_pw=owner_pwd,
                permissions=perm
            )
            doc.close()
            
            await PdfService._save_history(db, user_id, f"加密 PDF: {output_name}", output_name, "encrypt")

            return str(save_path)
        except Exception as e:
            logger.error(f"加密PDF失败: {e}")
            raise ValueError(f"加密失败: {str(e)}")

    @staticmethod
    async def decrypt_pdf(db: AsyncSession, user_id: int, request: PdfDecryptRequest) -> str:
        """解密 PDF"""
        try:
            fpath, old_name = await PdfService._resolve_file(db, user_id, request.file_id, request.path)
            
            output_name = request.output_name or f"decrypted_{old_name}"
            if not output_name.lower().endswith(".pdf"):
                output_name += ".pdf"
            
            # 使用密码打开加密的 PDF
            doc = fitz.open(fpath)
            if doc.is_encrypted:
                if not doc.authenticate(request.password):
                    raise ValueError("密码错误")
            
            save_path = PdfService._get_output_path(user_id, output_name)
            
            # 保存为未加密的 PDF
            doc.save(str(save_path))
            doc.close()
            
            await PdfService._save_history(db, user_id, f"解密 PDF: {output_name}", output_name, "decrypt")

            return str(save_path)
        except Exception as e:
            logger.error(f"解密PDF失败: {e}")
            raise ValueError(f"解密失败: {str(e)}")

    @staticmethod
    async def rotate_pdf(db: AsyncSession, user_id: int, request: PdfRotateRequest) -> str:
        """旋转 PDF 页面"""
        try:
            fpath, old_name = await PdfService._resolve_file(db, user_id, request.file_id, request.path)
            
            output_name = request.output_name or f"rotated_{old_name}"
            if not output_name.lower().endswith(".pdf"):
                output_name += ".pdf"
            
            doc = fitz.open(fpath)
            
            # 确定要旋转的页面
            if request.page_ranges:
                pages_to_rotate = PdfService._parse_page_ranges(request.page_ranges, doc.page_count)
            else:
                pages_to_rotate = list(range(doc.page_count))
            
            # 标准化旋转角度
            angle = request.angle % 360
            if angle not in [0, 90, 180, 270]:
                angle = 90  # 默认90度
            
            for page_num in pages_to_rotate:
                page = doc.load_page(page_num)
                page.set_rotation(page.rotation + angle)
            
            save_path = PdfService._get_output_path(user_id, output_name)
            doc.save(str(save_path))
            doc.close()
            
            await PdfService._save_history(db, user_id, f"旋转 PDF: {output_name}", output_name, "rotate")

            return str(save_path)
        except Exception as e:
            logger.error(f"旋转PDF失败: {e}")
            raise ValueError(f"旋转失败: {str(e)}")

    @staticmethod
    async def reorder_pages(db: AsyncSession, user_id: int, request: PdfReorderRequest) -> str:
        """重新排列 PDF 页面顺序"""
        try:
            fpath, old_name = await PdfService._resolve_file(db, user_id, request.file_id, request.path)
            
            output_name = request.output_name or f"reordered_{old_name}"
            if not output_name.lower().endswith(".pdf"):
                output_name += ".pdf"
            
            doc = fitz.open(fpath)
            
            # 验证页面顺序
            new_order = [p - 1 for p in request.page_order]  # 转换为 0-indexed
            if not all(0 <= p < doc.page_count for p in new_order):
                raise ValueError("页码顺序包含无效的页码")
            
            # 创建新文档并按新顺序插入页面
            result_doc = fitz.open()
            for page_num in new_order:
                result_doc.insert_pdf(doc, from_page=page_num, to_page=page_num)
            
            save_path = PdfService._get_output_path(user_id, output_name)
            result_doc.save(str(save_path))
            result_doc.close()
            doc.close()
            
            await PdfService._save_history(db, user_id, f"页面重排: {output_name}", output_name, "reorder")

            return str(save_path)
        except Exception as e:
            logger.error(f"重排PDF页面失败: {e}")
            raise ValueError(f"重排失败: {str(e)}")

    @staticmethod
    async def extract_pages(db: AsyncSession, user_id: int, request: PdfExtractPagesRequest) -> str:
        """提取 PDF 指定页面"""
        try:
            fpath, old_name = await PdfService._resolve_file(db, user_id, request.file_id, request.path)
            
            output_name = request.output_name or f"extracted_{old_name}"
            if not output_name.lower().endswith(".pdf"):
                output_name += ".pdf"
            
            doc = fitz.open(fpath)
            
            # 解析页码范围
            pages_to_extract = PdfService._parse_page_ranges(request.page_ranges, doc.page_count)
            
            if not pages_to_extract:
                raise ValueError("未选择有效的页码范围")
            
            # 使用 select 方法提取页面，保留元数据
            doc.select(pages_to_extract)
            
            save_path = PdfService._get_output_path(user_id, output_name)
            doc.save(str(save_path))
            doc.close()
            
            await PdfService._save_history(db, user_id, f"提取页面: {output_name}", output_name, "extract_pages")

            return str(save_path)
        except Exception as e:
            logger.error(f"提取PDF页面失败: {e}")
            raise ValueError(f"提取失败: {str(e)}")

    @staticmethod
    async def add_page_numbers(db: AsyncSession, user_id: int, request: PdfAddPageNumbersRequest) -> str:
        """为 PDF 添加页码"""
        try:
            fpath, old_name = await PdfService._resolve_file(db, user_id, request.file_id, request.path)
            
            output_name = request.output_name or f"numbered_{old_name}"
            if not output_name.lower().endswith(".pdf"):
                output_name += ".pdf"
            
            doc = fitz.open(fpath)
            total_pages = doc.page_count
            
            # 解析位置
            position_map = {
                "bottom-left": ("left", "bottom"),
                "bottom-center": ("center", "bottom"),
                "bottom-right": ("right", "bottom"),
                "top-left": ("left", "top"),
                "top-center": ("center", "top"),
                "top-right": ("right", "top"),
            }
            h_align, v_align = position_map.get(request.position, ("center", "bottom"))
            
            for i, page in enumerate(doc):
                rect = page.rect
                page_num = request.start_number + i
                
                # 格式化页码文本
                text = request.format.replace("{n}", str(page_num)).replace("{total}", str(total_pages))
                
                # 计算位置
                margin = 36  # 边距
                if h_align == "left":
                    x = margin
                elif h_align == "right":
                    x = rect.width - margin
                else:  # center
                    x = rect.width / 2
                
                if v_align == "top":
                    y = margin + request.fontsize
                else:  # bottom
                    y = rect.height - margin
                
                # 插入页码
                page.insert_text(
                    fitz.Point(x, y),
                    text,
                    fontsize=request.fontsize,
                    fontname="helv",
                    color=(0, 0, 0),
                    rotate=0
                )
            
            save_path = PdfService._get_output_path(user_id, output_name)
            doc.save(str(save_path))
            doc.close()
            
            await PdfService._save_history(db, user_id, f"添加页码: {output_name}", output_name, "add_page_numbers")

            return str(save_path)
        except Exception as e:
            logger.error(f"添加页码失败: {e}")
            raise ValueError(f"添加页码失败: {str(e)}")

    @staticmethod
    async def add_signature(db: AsyncSession, user_id: int, request: PdfSignRequest) -> str:
        """为 PDF 添加签名/印章图片"""
        try:
            fpath, old_name = await PdfService._resolve_file(db, user_id, request.file_id, request.path)
            img_path, _ = await PdfService._resolve_file(db, user_id, request.image_file_id, request.image_path)
            
            output_name = request.output_name or f"signed_{old_name}"
            if not output_name.lower().endswith(".pdf"):
                output_name += ".pdf"
            
            doc = fitz.open(fpath)
            
            # 获取目标页面
            page_num = request.page - 1  # 转换为 0-indexed
            if page_num < 0 or page_num >= doc.page_count:
                raise ValueError(f"页码超出范围 (1-{doc.page_count})")
            
            page = doc.load_page(page_num)
            rect = page.rect
            
            # 计算签名位置和大小
            x = rect.width * request.x / 100
            y = rect.height * request.y / 100
            width = rect.width * request.width / 100
            
            # 加载签名图片并计算高度
            img = fitz.open(img_path)
            if img.page_count > 0:
                img_page = img.load_page(0)
                aspect_ratio = img_page.rect.height / img_page.rect.width if img_page.rect.width > 0 else 1
            else:
                aspect_ratio = 1
            img.close()
            
            height = width * aspect_ratio
            
            # 创建签名区域
            sign_rect = fitz.Rect(x, y, x + width, y + height)
            
            # 插入图片
            page.insert_image(sign_rect, filename=img_path)
            
            save_path = PdfService._get_output_path(user_id, output_name)
            doc.save(str(save_path))
            doc.close()
            
            await PdfService._save_history(db, user_id, f"添加签名: {output_name}", output_name, "sign")

            return str(save_path)
        except Exception as e:
            logger.error(f"添加签名失败: {e}")
            raise ValueError(f"添加签名失败: {str(e)}")

    @staticmethod
    async def delete_pages(db: AsyncSession, user_id: int, file_id: Optional[int], page_ranges: str, output_name: Optional[str] = None, path: Optional[str] = None) -> str:
        """删除 PDF 指定页面"""
        try:
            fpath, old_name = await PdfService._resolve_file(db, user_id, file_id, path)
            
            output_name = output_name or f"deleted_{old_name}"
            if not output_name.lower().endswith(".pdf"):
                output_name += ".pdf"
            
            doc = fitz.open(fpath)
            
            # 解析要删除的页码
            pages_to_delete = set(PdfService._parse_page_ranges(page_ranges, doc.page_count))
            
            # 创建保留页面的列表
            pages_to_keep = [i for i in range(doc.page_count) if i not in pages_to_delete]
            
            if not pages_to_keep:
                raise ValueError("不能删除所有页面")
            
            # 使用 select 方法保留页面，保留元数据
            doc.select(pages_to_keep)
            
            save_path = PdfService._get_output_path(user_id, output_name)
            doc.save(str(save_path))
            doc.close()
            
            await PdfService._save_history(db, user_id, f"删除页面: {output_name}", output_name, "delete_pages")

            return str(save_path)
        except Exception as e:
            logger.error(f"删除PDF页面失败: {e}")
            raise ValueError(f"删除失败: {str(e)}")

    @staticmethod
    async def reverse_pages(db: AsyncSession, user_id: int, file_id: Optional[int], output_name: Optional[str] = None, path: Optional[str] = None) -> str:
        """反转 PDF 页面顺序"""
        try:
            fpath, old_name = await PdfService._resolve_file(db, user_id, file_id, path)
            
            output_name = output_name or f"reversed_{old_name}"
            if not output_name.lower().endswith(".pdf"):
                output_name += ".pdf"
            
            doc = fitz.open(fpath)
            
            # 创建新文档，按倒序插入页面
            result_doc = fitz.open()
            for page_num in range(doc.page_count - 1, -1, -1):
                result_doc.insert_pdf(doc, from_page=page_num, to_page=page_num)
            
            save_path = PdfService._get_output_path(user_id, output_name)
            result_doc.save(str(save_path))
            result_doc.close()
            doc.close()
            
            await PdfService._save_history(db, user_id, f"反转页面: {output_name}", output_name, "reverse")

            return str(save_path)
        except Exception as e:
            logger.error(f"反转PDF页面失败: {e}")
            raise ValueError(f"反转失败: {str(e)}")

    @staticmethod
    async def word_to_pdf(db: AsyncSession, user_id: int, request: PdfWordToPdfRequest) -> str:
        """Word 转 PDF (基础版: 仅提取文本内容)"""
        try:
            fpath, old_name = await PdfService._resolve_file(db, user_id, request.file_id, request.path)
            
            output_name = request.output_name or os.path.splitext(old_name)[0] + ".pdf"
            if not output_name.lower().endswith(".pdf"):
                output_name += ".pdf"
            
            save_path = PdfService._get_output_path(user_id, output_name)
            
            # 使用 python-docx 读取
            try:
                from docx import Document
                doc = Document(fpath)
            except ImportError:
                raise ValueError("服务器缺少 python-docx 库")
            
            # 创建 PDF
            pdf = fitz.open()
            page = pdf.new_page()
            margin = 50
            y = margin
            line_height = 14
            
            # 使用内置字体，不支持复杂中文排版是已知限制
            fontname = "china-s" 
            
            for para in doc.paragraphs:
                text = para.text
                if not text.strip():
                    y += line_height / 2
                    continue
                
                # 简单的自动换行处理
                page.insert_text((margin, y), text, fontsize=11, fontname=fontname, encoding=0)
                y += line_height * 1.5
                
                if y > page.rect.height - margin:
                    page = pdf.new_page()
                    y = margin

            pdf.save(str(save_path))
            pdf.close()
            
            await PdfService._save_history(db, user_id, f"Word转PDF: {output_name}", output_name, "word_to_pdf")

            return str(save_path)
        except Exception as e:
            logger.error(f"Word转PDF失败: {e}")
            raise ValueError(f"转换失败: {str(e)}")

    @staticmethod
    async def excel_to_pdf(db: AsyncSession, user_id: int, request: PdfExcelToPdfRequest) -> str:
        """Excel 转 PDF (基础版: 表格转文本)"""
        try:
            fpath, old_name = await PdfService._resolve_file(db, user_id, request.file_id, request.path)
            
            output_name = request.output_name or os.path.splitext(old_name)[0] + ".pdf"
            if not output_name.lower().endswith(".pdf"):
                output_name += ".pdf"
            
            save_path = PdfService._get_output_path(user_id, output_name)
            
            try:
                import openpyxl
            except ImportError:
                raise ValueError("服务器缺少 openpyxl 库")
            
            wb = openpyxl.load_workbook(fpath, data_only=True)
            
            pdf = fitz.open()
            fontname = "china-s"
            
            for sheetname in wb.sheetnames:
                ws = wb[sheetname]
                page = pdf.new_page()
                page.insert_text((50, 40), f"Sheet: {sheetname}", fontsize=14, fontname=fontname)
                y = 70
                
                for row in ws.iter_rows(values_only=True):
                    # 简单的 Tab 分隔
                    text = "  ".join([str(c) if c is not None else "" for c in row])
                    if text.strip():
                        page.insert_text((50, y), text, fontsize=10, fontname=fontname)
                        y += 12
                    
                    if y > page.rect.height - 50:
                        page = pdf.new_page()
                        y = 50
            
            pdf.save(str(save_path))
            pdf.close()
            wb.close()
            
            await PdfService._save_history(db, user_id, f"Excel转PDF: {output_name}", output_name, "excel_to_pdf")

            return str(save_path)
        except Exception as e:
            logger.error(f"Excel转PDF失败: {e}")
            raise ValueError(f"转换失败: {str(e)}")

    @staticmethod
    async def save_extracted_text(db: AsyncSession, user_id: int, request: PdfSaveTextRequest) -> str:
        """保存提取的文本到文件"""
        try:
            output_name = request.filename
            if not output_name.lower().endswith('.txt'): 
                output_name += '.txt'
            
            save_path = PdfService._get_output_path(user_id, output_name)
            
            with open(save_path, 'w', encoding='utf-8') as f:
                f.write(request.text)
                
            await PdfService._save_history(db, user_id, f"提取文本: {output_name}", output_name, "extract_text")

            return str(save_path)
        except Exception as e:
            logger.error(f"保存文本失败: {e}")
            raise ValueError(f"保存失败: {str(e)}")
